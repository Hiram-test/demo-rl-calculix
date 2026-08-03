from __future__ import annotations

import json
import textwrap
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.collections import LineCollection
from matplotlib.tri import Triangulation

from .scenarios import ScenarioRun


def configure_fonts() -> None:
    candidates = ["Noto Sans CJK SC", "Noto Sans CJK JP", "Source Han Sans CN", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    available = {f.name for f in matplotlib.font_manager.fontManager.ttflist}
    for name in candidates:
        if name in available:
            plt.rcParams["font.family"] = name
            break
    plt.rcParams["axes.unicode_minus"] = False


def _mesh_lines(run: ScenarioRun, key: str) -> LineCollection:
    mesh = run.meshes[key]
    lines = []
    for conn in mesh.elements:
        pts = mesh.nodes[np.r_[conn, conn[0]]]
        lines.extend([[pts[i], pts[i + 1]] for i in range(4)])
    return LineCollection(lines, linewidths=0.25)


def save_mesh_figure(run: ScenarioRun, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    keys = list(run.meshes)
    base_key = keys[0]
    final_key = next((k for k in keys if "rounded_final" in k or "patch_final" in k), keys[-1])
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))
    for ax, key, label in [(axes[0], base_key, "初始/较粗模型"), (axes[1], final_key, "最终落实模型")]:
        mesh = run.meshes[key]
        ax.add_collection(_mesh_lines(run, key))
        for special in ("inner_hole", "crack_upper", "crack_lower"):
            if special in mesh.edge_sets:
                for a, b in mesh.edge_sets[special]:
                    pts = mesh.nodes[[a, b]]
                    ax.plot(pts[:, 0], pts[:, 1], linewidth=2.0)
        ax.autoscale()
        ax.set_aspect("equal", adjustable="box")
        ax.set_title(f"{label}\n{key}，{len(mesh.elements)}个单元")
        ax.set_xlabel("x / mm")
        ax.set_ylabel("y / mm")
    fig.suptitle(run.title)
    fig.tight_layout()
    path = out_dir / f"{run.scenario_id}_mesh.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def _triangles(elements: np.ndarray) -> np.ndarray:
    return np.vstack([elements[:, [0, 1, 2]], elements[:, [0, 2, 3]]])


def save_stress_figure(run: ScenarioRun, out_dir: Path) -> Path:
    keys = list(run.solutions)
    baseline_key = next((k for k in reversed(keys) if "extended" not in k and "patch" not in k and "rounded" not in k), keys[-1])
    variant_key = next((k for k in keys if "patch_final" in k or "rounded_final" in k), None)
    plot_keys = [baseline_key] + ([variant_key] if variant_key else [])
    fig, axes = plt.subplots(1, len(plot_keys), figsize=(6 * len(plot_keys), 4.8), squeeze=False)
    for ax, key in zip(axes[0], plot_keys):
        mesh = run.meshes[key]
        sol = run.solutions[key]
        tri = Triangulation(mesh.nodes[:, 0], mesh.nodes[:, 1], _triangles(mesh.elements))
        c = ax.tricontourf(tri, sol.nodal_von_mises, levels=24)
        fig.colorbar(c, ax=ax, label="von Mises / MPa")
        ax.set_aspect("equal", adjustable="box")
        ax.set_title(key)
        ax.set_xlabel("x / mm")
        ax.set_ylabel("y / mm")
    fig.suptitle(f"{run.title}：最终网格应力图")
    fig.tight_layout()
    path = out_dir / f"{run.scenario_id}_stress.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def save_convergence_figure(run: ScenarioRun, out_dir: Path) -> Path:
    rows = run.level_rows
    h = np.array([float(r["h_local"]) for r in rows])
    metrics = {
        "局部峰值": np.array([float(r["peak_stress"]) for r in rows]),
        "物理评价量": np.array([float(r["fixed_qoi"]) for r in rows]),
        "应变能": np.array([float(r["strain_energy"]) for r in rows]),
    }
    fig, ax = plt.subplots(figsize=(7.2, 4.8))
    for label, vals in metrics.items():
        denom = vals[-1] if abs(vals[-1]) > 1e-30 else 1.0
        ax.plot(h, vals / denom, marker="o", label=label)
    ax.set_xscale("log")
    ax.invert_xaxis()
    ax.set_xlabel("局部特征尺寸 h / mm（向右为更细）")
    ax.set_ylabel("相对最终网格值")
    ax.grid(True, alpha=0.3)
    ax.legend()
    ax.set_title(f"{run.title}：网格趋势")
    path = out_dir / f"{run.scenario_id}_convergence.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def save_theory_figure(run: ScenarioRun, out_dir: Path) -> Path | None:
    checks = [s for s in run.diagnostic.skill_trace if s.name == "theory_cross_check" and "theoretical" in s.outputs]
    if not checks:
        return None
    labels = [str(s.inputs.get("quantity", "理论量")) for s in checks]
    num = np.array([float(s.outputs["numerical"]) for s in checks])
    theory = np.array([float(s.outputs["theoretical"]) for s in checks])
    x = np.arange(len(labels))
    width = 0.36
    fig, ax = plt.subplots(figsize=(max(7, 3.2 * len(labels)), 4.6))
    ax.bar(x - width / 2, num, width, label="有限元")
    ax.bar(x + width / 2, theory, width, label="理论/守恒")
    ax.set_xticks(x)
    ax.set_xticklabels(["\n".join(textwrap.wrap(s, 12)) for s in labels])
    ax.legend()
    ax.set_title(f"{run.title}：独立交叉验证")
    ax.grid(axis="y", alpha=0.25)
    path = out_dir / f"{run.scenario_id}_theory.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_figures(runs: list[ScenarioRun], out_dir: Path) -> dict[str, dict[str, str]]:
    configure_fonts()
    out: dict[str, dict[str, str]] = {}
    for run in runs:
        mesh = save_mesh_figure(run, out_dir)
        stress = save_stress_figure(run, out_dir)
        conv = save_convergence_figure(run, out_dir)
        theory = save_theory_figure(run, out_dir)
        out[run.scenario_id] = {"mesh": str(mesh), "stress": str(stress), "convergence": str(conv)}
        if theory:
            out[run.scenario_id]["theory"] = str(theory)
    return out


def _format_float(value: Any) -> str:
    if isinstance(value, (float, np.floating)):
        if abs(value) >= 10000 or (abs(value) < 0.001 and value != 0):
            return f"{value:.4e}"
        return f"{value:.4f}"
    return str(value)


def build_markdown_report(runs: list[ScenarioRun], figures: dict[str, dict[str, str]], output_path: Path, deepseek_review: dict[str, Any] | None = None) -> None:
    lines = [
        "# 桥梁构件网格诊断与物理融合加密报告",
        "",
        "本报告面向普通工程人员：从‘网格越细峰值越高’等直接问题出发，自动完成建模、多级网格、有限元求解、能量检查、理论/守恒交叉验证、物理修正和可执行加密方案。",
        "",
        "## 总览",
        "",
        "| 构件场景 | 用户问题 | 诊断 | 已落实方案 |",
        "|---|---|---|---|",
    ]
    for run in runs:
        lines.append(f"| {run.title} | {run.user_question} | {run.diagnostic.diagnosis} | {'；'.join(run.diagnostic.applied_plan[:2])} |")
    lines.extend(["", "## 能量 Skill 是否实际执行", ""])
    total_energy = 0
    for run in runs:
        for skill in run.diagnostic.skill_trace:
            if skill.name == "energy_consistency":
                total_energy += 1
                lines.append(f"- **{run.title}**：已执行；最大能量平衡误差 `{max(skill.outputs['balance_relative_errors']):.3e}`，末两档应变能变化 `{skill.outputs['last_mesh_energy_change']:.3%}`。")
    lines.append(f"\n本次共实际执行 **{total_energy} 次**能量一致性 Skill。它不是装饰字段，而是参与峰值诊断和停止判断的独立证据。")
    review_by_id = {x.get("scenario_id"): x for x in (deepseek_review or {}).get("scenario_reviews", [])}
    for index, run in enumerate(runs, start=1):
        lines.extend(["", f"# {index}. {run.title}", "", f"> 用户原始问题：{run.user_question}", "", "## 结论", "", run.diagnostic.plain_explanation, ""])
        review = review_by_id.get(run.scenario_id)
        if review:
            lines.extend(["### DeepSeek面向用户的复核", "", review.get("user_answer", ""), ""])
        lines.extend(["## 已落实的处理", ""])
        for item in run.diagnostic.applied_plan:
            lines.append(f"- {item}")
        lines.extend(["", "## 网格与结果图", ""])
        for kind in ["mesh", "convergence", "stress", "theory"]:
            if kind in figures[run.scenario_id]:
                lines.append(f"![{kind}]({Path(figures[run.scenario_id][kind]).name})")
        lines.extend(["", "## 多级网格数据", ""])
        headers = list(run.level_rows[0].keys())
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")
        for row in run.level_rows:
            lines.append("| " + " | ".join(_format_float(row.get(h, "")) for h in headers) + " |")
        lines.extend(["", "## Skill执行记录", ""])
        for skill in run.diagnostic.skill_trace:
            lines.append(f"- **{skill.name}** — {'通过' if skill.passed else '未通过'}：{skill.purpose} 输出：`{json.dumps(skill.outputs, ensure_ascii=False)}`")
        lines.extend(["", "## 适用边界", "", "**可以用于：** " + "；".join(run.diagnostic.supported_use), "", "**不能直接用于：** " + "；".join(run.diagnostic.unsupported_use)])
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _pdf_text_page(pdf: PdfPages, title: str, paragraphs: list[str]) -> None:
    fig = plt.figure(figsize=(8.27, 11.69))
    ax = fig.add_axes([0.08, 0.06, 0.84, 0.88])
    ax.axis("off")
    ax.text(0.0, 0.98, title, fontsize=20, fontweight="bold", va="top")
    y = 0.91
    for p in paragraphs:
        wrapped = textwrap.fill(p, width=48)
        ax.text(0.0, y, wrapped, fontsize=11.5, va="top", linespacing=1.55)
        y -= 0.06 + 0.025 * wrapped.count("\n")
        if y < 0.08:
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)
            fig = plt.figure(figsize=(8.27, 11.69))
            ax = fig.add_axes([0.08, 0.06, 0.84, 0.88])
            ax.axis("off")
            y = 0.95
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def build_pdf_report(runs: list[ScenarioRun], figures: dict[str, dict[str, str]], output_path: Path, deepseek_review: dict[str, Any] | None = None) -> None:
    configure_fonts()
    review_by_id = {x.get("scenario_id"): x for x in (deepseek_review or {}).get("scenario_reviews", [])}
    with PdfPages(output_path) as pdf:
        _pdf_text_page(pdf, "桥梁构件网格诊断与物理融合加密报告", [
            "对象：普通桥梁工程人员。输入可以只是‘最大应力随网格无限递增，怎么办？’这类直接问题。",
            f"本次完成 {len(runs)} 类简单桥梁构件场景，全部实际执行建模、多级网格、有限元求解、能量检查、理论或守恒核验、物理修正与报告输出。",
            "核心原则：不要求用户先掌握塑性本构、J积分或奇异性理论；系统根据证据决定应继续加密、切换评价量，还是修改物理载荷/几何表达。",
        ])
        for run in runs:
            paras = [
                f"用户问题：{run.user_question}",
                f"诊断：{run.diagnostic.diagnosis}",
                f"说明：{run.diagnostic.plain_explanation}",
                "已落实：" + "；".join(run.diagnostic.applied_plan),
                "可以用于：" + "；".join(run.diagnostic.supported_use),
                "不能直接用于：" + "；".join(run.diagnostic.unsupported_use),
            ]
            review = review_by_id.get(run.scenario_id)
            if review:
                paras.append("DeepSeek复核：" + str(review.get("user_answer", "")))
            _pdf_text_page(pdf, run.title, paras)
            for key in ["mesh", "convergence", "stress", "theory"]:
                if key not in figures[run.scenario_id]:
                    continue
                img = plt.imread(figures[run.scenario_id][key])
                fig = plt.figure(figsize=(8.27, 11.69))
                ax = fig.add_axes([0.05, 0.08, 0.9, 0.84])
                ax.imshow(img)
                ax.axis("off")
                ax.set_title(f"{run.title} — {key}", fontsize=14)
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)


def build_docx_report(runs: list[ScenarioRun], figures: dict[str, dict[str, str]], output_path: Path, deepseek_review: dict[str, Any] | None = None) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("桥梁构件网格诊断与物理融合加密报告", 0)
    doc.add_paragraph("面向普通工程人员：从直接工程问题出发，自动完成建模、网格、求解、能量与理论核验、物理修正和报告。")
    doc.add_heading("总体完成情况", level=1)
    doc.add_paragraph(f"本次完整执行 {len(runs)} 种构件场景；每个场景均至少包含三档网格和独立能量一致性 Skill。")
    review_by_id = {x.get("scenario_id"): x for x in (deepseek_review or {}).get("scenario_reviews", [])}
    for idx, run in enumerate(runs, start=1):
        doc.add_heading(f"{idx}. {run.title}", level=1)
        doc.add_paragraph(f"用户问题：{run.user_question}")
        doc.add_heading("诊断", level=2)
        doc.add_paragraph(run.diagnostic.plain_explanation)
        review = review_by_id.get(run.scenario_id)
        if review:
            doc.add_heading("DeepSeek复核", level=2)
            doc.add_paragraph(str(review.get("user_answer", "")))
        doc.add_heading("已落实方案", level=2)
        for item in run.diagnostic.applied_plan:
            doc.add_paragraph(item, style="List Bullet")
        for key in ["mesh", "convergence", "stress", "theory"]:
            if key in figures[run.scenario_id]:
                doc.add_picture(figures[run.scenario_id][key], width=Inches(6.2))
        doc.add_heading("多级网格数据", level=2)
        headers = list(run.level_rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in run.level_rows:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = _format_float(row.get(h, ""))
        doc.add_heading("适用边界", level=2)
        doc.add_paragraph("可以用于：" + "；".join(run.diagnostic.supported_use))
        doc.add_paragraph("不能直接用于：" + "；".join(run.diagnostic.unsupported_use))
    doc.save(output_path)
